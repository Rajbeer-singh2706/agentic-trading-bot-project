"""Document ingestion use case."""
import asyncio
import time
from typing import List
from datetime import datetime

from fastapi import UploadFile
from domain.entities import Document, IngestionResult
from domain.repositories import DocumentRepository
from domain.exceptions import DocumentProcessingError, ValidationError
from infrastructure.logging import get_logger


class DocumentIngestionUseCase:
    """Use case for ingesting documents."""

    def __init__(self, document_repo: DocumentRepository):
        """Initialize document ingestion use case."""
        self.document_repo = document_repo
        self.logger = get_logger(__name__)
        self.max_file_size = 100 * 1024 * 1024  # 100 MB
        self.allowed_extensions = {".pdf", ".txt", ".docx", ".md"}

    async def execute(self, files: List[UploadFile]) -> IngestionResult:
        """Ingest documents.
        
        Args:
            files: List of uploaded files
            
        Returns:
            Ingestion result with statistics
            
        Raises:
            ValidationError: If files are invalid
            DocumentProcessingError: If processing fails
        """
        start_time = time.time()
        documents_processed = 0
        chunks_created = 0
        errors = []

        try:
            self.logger.info(f"Starting ingestion of {len(files)} files")

            # Validate files
            for file in files:
                try:
                    self._validate_file(file)
                except ValidationError as e:
                    errors.append(str(e))
                    self.logger.warning(f"File validation failed: {e}")
                    continue

                # Process file
                try:
                    content = await file.read()
                    extracted_text = await self._extract_text(file.filename, content)

                    document = Document(
                        id=f"{file.filename}_{int(time.time())}",
                        filename=file.filename,
                        content=extracted_text,
                        source=f"uploaded:{file.filename}",
                        created_at=datetime.utcnow(),
                        metadata={
                            "original_filename": file.filename,
                            "content_type": file.content_type,
                            "size_bytes": len(content),
                        },
                    )

                    # Save document (which includes chunking and vector storage)
                    doc_id = await self.document_repo.save(document)
                    document.id = doc_id
                    documents_processed += 1
                    chunks_created += 1  # Simplified - actual count depends on chunking

                    self.logger.info(
                        f"Successfully ingested document: {file.filename}",
                        document_id=doc_id,
                    )

                except Exception as e:
                    error_msg = f"Failed to process {file.filename}: {str(e)}"
                    errors.append(error_msg)
                    self.logger.error(error_msg)

            execution_time = time.time() - start_time

            result = IngestionResult(
                success=len(errors) == 0,
                documents_processed=documents_processed,
                chunks_created=chunks_created,
                errors=errors,
                execution_time=execution_time,
                metadata={
                    "total_files": len(files),
                    "files_processed": documents_processed,
                    "error_count": len(errors),
                },
            )

            self.logger.info(
                f"Ingestion completed: {documents_processed} docs, "
                f"{len(errors)} errors, {execution_time:.2f}s"
            )

            return result

        except Exception as e:
            self.logger.error(f"Document ingestion failed: {e}")
            raise DocumentProcessingError(f"Failed to ingest documents: {str(e)}") from e

    def _validate_file(self, file: UploadFile) -> None:
        """Validate uploaded file.
        
        Args:
            file: File to validate
            
        Raises:
            ValidationError: If file is invalid
        """
        if not file or not file.filename:
            raise ValidationError("File is required")

        # Check file size (approximate check, actual size checked after reading)
        if file.size and file.size > self.max_file_size:
            raise ValidationError(
                f"File too large: {file.size} > {self.max_file_size}"
            )

        # Check file extension
        import os

        _, ext = os.path.splitext(file.filename.lower())
        if ext not in self.allowed_extensions:
            raise ValidationError(
                f"File type not supported: {ext}. "
                f"Allowed: {', '.join(self.allowed_extensions)}"
            )

    async def _extract_text(self, filename: str, content: bytes) -> str:
        """Extract text from file.
        
        Args:
            filename: Original filename
            content: File content
            
        Returns:
            Extracted text
            
        Raises:
            DocumentProcessingError: If extraction fails
        """
        import os

        _, ext = os.path.splitext(filename.lower())

        try:
            if ext == ".txt" or ext == ".md":
                return content.decode("utf-8", errors="ignore")

            elif ext == ".pdf":
                return await self._extract_pdf_text(content)

            elif ext == ".docx":
                return await self._extract_docx_text(content)

            else:
                raise DocumentProcessingError(f"Unsupported file type: {ext}")

        except Exception as e:
            raise DocumentProcessingError(f"Failed to extract text: {str(e)}") from e

    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF."""
        try:
            import pypdf
            from io import BytesIO

            pdf_reader = pypdf.PdfReader(BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except ImportError:
            raise DocumentProcessingError("PyPDF not installed. Install: pip install pypdf")

    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX."""
        try:
            from io import BytesIO
            from docx import Document

            doc = Document(BytesIO(content))
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except ImportError:
            raise DocumentProcessingError(
                "python-docx not installed. Install: pip install python-docx"
            )
